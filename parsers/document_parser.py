"""
parsers/pdf_parser.py — Extract text from PDF resumes using pdfplumber.
parsers/docx_parser.py — Extract text from DOCX resumes using python-docx.
parsers/text_handler.py — Handle plain text and normalize file inputs.

All parsers return a normalized string ready for analysis and prompting.
"""

import io
import os
from pathlib import Path
from typing import Union

from utils.logger import get_logger
from utils.helpers import clean_text

logger = get_logger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# PDF PARSER
# ═══════════════════════════════════════════════════════════════════════════════

class PDFParser:
    """
    Extracts text from PDF files using pdfplumber.
    Falls back to PyMuPDF (fitz) if pdfplumber fails.
    """

    def parse_file(self, file_path: Union[str, Path]) -> str:
        """Extract text from a PDF file path."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        logger.info("Parsing PDF: %s", file_path.name)
        return self._extract_with_pdfplumber(file_path)

    def parse_bytes(self, content: bytes) -> str:
        """Extract text from PDF bytes (e.g., from an uploaded file)."""
        logger.info("Parsing PDF from bytes (%d bytes)", len(content))
        return self._extract_with_pdfplumber(io.BytesIO(content))

    def _extract_with_pdfplumber(self, source: Union[Path, io.BytesIO]) -> str:
        try:
            import pdfplumber
            with pdfplumber.open(source) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                    else:
                        logger.warning("Page %d yielded no text (may be image-based)", i + 1)

            if not pages:
                raise ValueError("No text could be extracted from PDF")

            raw = "\n\n".join(pages)
            return clean_text(raw)

        except ImportError:
            logger.warning("pdfplumber not available, trying PyMuPDF")
            return self._extract_with_pymupdf(source)

    def _extract_with_pymupdf(self, source: Union[Path, io.BytesIO]) -> str:
        try:
            import fitz  # PyMuPDF
            if isinstance(source, io.BytesIO):
                doc = fitz.open(stream=source.read(), filetype="pdf")
            else:
                doc = fitz.open(str(source))

            pages = [page.get_text() for page in doc]
            doc.close()
            return clean_text("\n\n".join(pages))

        except ImportError:
            raise ImportError(
                "Neither pdfplumber nor PyMuPDF is installed. "
                "Install with: pip install pdfplumber or pip install pymupdf"
            )


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX PARSER
# ═══════════════════════════════════════════════════════════════════════════════

class DOCXParser:
    """Extracts text from .docx files using python-docx."""

    def parse_file(self, file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        logger.info("Parsing DOCX: %s", file_path.name)
        return self._extract(str(file_path))

    def parse_bytes(self, content: bytes) -> str:
        logger.info("Parsing DOCX from bytes (%d bytes)", len(content))
        return self._extract(io.BytesIO(content))

    def _extract(self, source) -> str:
        try:
            from docx import Document
            doc = Document(source)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables (skills tables are common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            raw = "\n".join(paragraphs)
            return clean_text(raw)
        except ImportError:
            raise ImportError(
                "python-docx is not installed. Install with: pip install python-docx"
            )


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT HANDLER
# ═══════════════════════════════════════════════════════════════════════════════

class TextHandler:
    """Handles plain text inputs — normalizes and validates."""

    def process(self, text: str) -> str:
        if not text or not text.strip():
            raise ValueError("Input text is empty")
        return clean_text(text)

    def read_file(self, file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return self.process(f.read())


# ═══════════════════════════════════════════════════════════════════════════════
# UNIFIED DOCUMENT PARSER
# ═══════════════════════════════════════════════════════════════════════════════

class DocumentParser:
    """
    Unified entry point for all document types.
    Auto-detects format from file extension or MIME type.
    """

    def __init__(self):
        self.pdf_parser = PDFParser()
        self.docx_parser = DOCXParser()
        self.text_handler = TextHandler()

    def parse_upload(self, filename: str, content: bytes) -> str:
        """
        Parse an uploaded file based on its extension.

        Args:
            filename: Original filename (used to determine type).
            content: Raw bytes of the file.

        Returns:
            Extracted and cleaned text string.
        """
        ext = Path(filename).suffix.lower().lstrip(".")

        if ext == "pdf":
            return self.pdf_parser.parse_bytes(content)
        elif ext in ("docx", "doc"):
            return self.docx_parser.parse_bytes(content)
        elif ext in ("txt", "text", "md"):
            return self.text_handler.process(content.decode("utf-8", errors="replace"))
        else:
            raise ValueError(
                f"Unsupported file type: .{ext}. "
                "Supported: pdf, docx, txt"
            )

    def parse_text_input(self, text: str) -> str:
        """Parse a raw text string input."""
        return self.text_handler.process(text)


# ─── Singleton ────────────────────────────────────────────────────────────────
document_parser = DocumentParser()
